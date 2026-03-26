from flask import Blueprint, send_file, jsonify, request
from flask_login import login_required
from models.base import get_db
from sqlalchemy import text
from docxtpl import DocxTemplate
from datetime import datetime
import os
import tempfile

# Docx imports for template creation
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_bp = Blueprint('documents', __name__)

@doc_bp.route('/projects/<int:project_id>/generate-report', methods=['POST'])
@login_required
def generate_board_report(project_id):
    """Generate professional Word document report"""
    
    data = request.get_json() or {}
    doc_type = data.get('type', 'board_analysis')
    
    db = next(get_db())
    
    # Get project data
    project_result = db.execute(text("""
        SELECT client_name, board_name, company_ticker, industry, status
        FROM projects
        WHERE id = :pid AND deleted_at IS NULL
    """), {"pid": project_id})
    
    project_row = project_result.fetchone()
    if not project_row:
        return jsonify({'error': 'Project not found'}), 404
    
    project = {
        'client_name': project_row[0],
        'board_name': project_row[1],
        'company_ticker': project_row[2],
        'industry': project_row[3],
        'status': project_row[4]
    }
    
    # Get board members
    board_result = db.execute(text("""
        SELECT name, organization, position, matrix_data
        FROM board_members
        WHERE project_id = :pid AND deleted_at IS NULL
    """), {"pid": project_id})
    
    members = []
    female_count = 0
    diverse_count = 0
    ai_ready_count = 0
    
    ai_keywords = ['AI', 'ML', 'Digital Transformation', 'Technology', 'Cybersecurity', 
                   'Data', 'Cloud', 'Software', 'Innovation', 'Tech', 'Machine Learning']
    
    for row in board_result:
        matrix_data = row[3] or {}
        gender = matrix_data.get('demographics', {}).get('gender', '')
        race = matrix_data.get('demographics', {}).get('race_ethnicity', '')
        
        background = matrix_data.get('background', {})
        bio = background.get('bio', '') if isinstance(background, dict) else ''
        
        is_ai_ready = any(keyword.lower() in bio.lower() for keyword in ai_keywords) if bio else False
        if is_ai_ready:
            ai_ready_count += 1
        
        if gender == 'Female':
            female_count += 1
        if race and race != 'Caucasian':
            diverse_count += 1
        
        members.append({
            'name': row[0],
            'organization': row[1],
            'position': row[2],
            'gender': gender,
            'race': race,
            'ai_ready': is_ai_ready,
            'bio': bio
        })
    
    total_members = len(members)
    
    # Get gaps
    gaps_result = db.execute(text("""
        SELECT gap_title, gap_description, priority, priority_score
        FROM gap_analysis
        WHERE project_id = :pid AND deleted_at IS NULL
        ORDER BY priority_score DESC
    """), {"pid": project_id})
    
    gaps = []
    for row in gaps_result:
        gaps.append({
            'title': row[0],
            'description': row[1],
            'priority': row[2].upper() if row[2] else 'MEDIUM',
            'score': row[3] or 50
        })
    
    # Calculate metrics
    female_percentage = round((female_count / total_members * 100), 1) if total_members > 0 else 0
    diverse_percentage = round((diverse_count / total_members * 100), 1) if total_members > 0 else 0
    ai_ready_percentage = round((ai_ready_count / total_members * 100), 1) if total_members > 0 else 0
    
    # AI Readiness Grade
    if ai_ready_percentage >= 40:
        ai_grade, ai_rating = 'A', 'Excellent'
    elif ai_ready_percentage >= 30:
        ai_grade, ai_rating = 'B', 'Good'
    elif ai_ready_percentage >= 20:
        ai_grade, ai_rating = 'C', 'Fair'
    elif ai_ready_percentage >= 10:
        ai_grade, ai_rating = 'D', 'Poor'
    else:
        ai_grade, ai_rating = 'F', 'Critical Gap'
    
    # Defense Score
    defense_score = 100
    if female_percentage < 30:
        defense_score -= 10
    if diverse_percentage < 25:
        defense_score -= 10
    defense_score = max(0, defense_score)
    
    if defense_score >= 80:
        defense_rating, defense_color = 'Strong', 'Green'
    elif defense_score >= 60:
        defense_rating, defense_color = 'Moderate', 'Yellow'
    else:
        defense_rating, defense_color = 'Vulnerable', 'Red'
    
    # Context
    context = {
        'client_name': project['client_name'],
        'board_name': project['board_name'],
        'company_ticker': project['company_ticker'] or 'N/A',
        'industry': project['industry'] or 'N/A',
        'report_date': datetime.now().strftime('%B %d, %Y'),
        'total_members': total_members,
        'female_count': female_count,
        'male_count': total_members - female_count,
        'female_percentage': female_percentage,
        'diverse_count': diverse_count,
        'diverse_percentage': diverse_percentage,
        'sp500_female': 32,
        'sp500_diverse': 28,
        'members': members,
        'gaps': gaps,
        'critical_gaps': len([g for g in gaps if g['priority'] == 'CRITICAL']),
        'high_gaps': len([g for g in gaps if g['priority'] == 'HIGH']),
        'overall_diversity_score': round((female_percentage + diverse_percentage) / 2, 1),
        'gender_vs_benchmark': round(female_percentage - 32, 1),
        'diversity_vs_benchmark': round(diverse_percentage - 28, 1),
        'ai_ready_count': ai_ready_count,
        'ai_ready_percentage': ai_ready_percentage,
        'ai_grade': ai_grade,
        'ai_rating': ai_rating,
        'ai_blind_members': total_members - ai_ready_count,
        'defense_score': defense_score,
        'defense_rating': defense_rating,
        'defense_color': defense_color
    }
    
    template_path = get_template_path(doc_type)
    
    if not os.path.exists(template_path):
        os.makedirs(os.path.dirname(template_path), exist_ok=True)
        create_template(template_path, doc_type)
    
    doc = DocxTemplate(template_path)
    doc.render(context)
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    filename = get_filename(project['client_name'], doc_type)
    
    return send_file(
        temp_file.name,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def get_template_path(doc_type):
    base_path = os.path.join(os.path.dirname(__file__), '..', 'templates')
    templates = {
        'board_analysis': 'board_analysis_template.docx',
        'ai_readiness': 'ai_readiness_template.docx',
        'activist_defense': 'activist_defense_template.docx',
        'strategic_alignment': 'strategic_alignment_template.docx',
        'interlock_map': 'interlock_map_template.docx',
        'diversity_scorecard': 'diversity_scorecard_template.docx',
        'executive_brief': 'executive_brief_template.docx',
        'gap_summary': 'gap_summary_template.docx',
        'candidate_dossier': 'candidate_dossier_template.docx'
    }
    return os.path.join(base_path, templates.get(doc_type, 'board_analysis_template.docx'))

def get_filename(client_name, doc_type):
    clean_name = client_name.replace(' ', '_')
    date_str = datetime.now().strftime('%Y%m%d')
    filenames = {
        'board_analysis': f'{clean_name}_Board_Analysis_{date_str}.docx',
        'ai_readiness': f'{clean_name}_AI_Readiness_{date_str}.docx',
        'activist_defense': f'{clean_name}_Activist_Defense_{date_str}.docx',
        'strategic_alignment': f'{clean_name}_Strategic_Alignment_{date_str}.docx',
        'interlock_map': f'{clean_name}_Interlock_Map_{date_str}.docx',
        'diversity_scorecard': f'{clean_name}_Diversity_Scorecard_{date_str}.docx',
        'executive_brief': f'{clean_name}_Executive_Brief_{date_str}.docx',
        'gap_summary': f'{clean_name}_Gap_Summary_{date_str}.docx',
        'candidate_dossier': f'{clean_name}_Candidate_Dossier_{date_str}.docx'
    }
    return filenames.get(doc_type, f'{clean_name}_Report_{date_str}.docx')

def create_template(template_path, doc_type):
    doc = Document()
    
    if doc_type == 'ai_readiness':
        create_ai_readiness_template(doc)
    elif doc_type == 'activist_defense':
        create_activist_defense_template(doc)
    elif doc_type == 'board_analysis':
        create_board_analysis_template(doc)
    else:
        create_simple_template(doc, doc_type)
    
    doc.save(template_path)

def create_ai_readiness_template(doc):
    title = doc.add_heading('AI Readiness Scorecard', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Client: {{ client_name }}')
    doc.add_paragraph('Board: {{ board_name }}')
    doc.add_paragraph('Date: {{ report_date }}')
    doc.add_paragraph()
    
    doc.add_heading('AI Readiness Grade', 1)
    doc.add_paragraph('Overall Grade: {{ ai_grade }} ({{ ai_rating }})')
    doc.add_paragraph('AI-Ready Directors: {{ ai_ready_count }} of {{ total_members }} ({{ ai_ready_percentage }}%)')
    doc.add_paragraph('AI-Blind Directors: {{ ai_blind_members }}')
    doc.add_paragraph()
    
    doc.add_heading('Analysis', 1)
    doc.add_paragraph(
        'This scorecard analyzes board members for AI/ML expertise based on '
        'keywords: AI, Machine Learning, Digital Transformation, Technology, Data, Cloud.'
    )

def create_activist_defense_template(doc):
    title = doc.add_heading('Activist Defense Audit', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Client: {{ client_name }}')
    doc.add_paragraph('Date: {{ report_date }}')
    doc.add_paragraph()
    
    doc.add_heading('Defense Score', 1)
    doc.add_paragraph('Overall Score: {{ defense_score }}/100 ({{ defense_rating }})')
    doc.add_paragraph()
    
    doc.add_heading('Board Composition', 1)
    doc.add_paragraph('Total Members: {{ total_members }}')
    doc.add_paragraph('Gender Diversity: {{ female_percentage }}%')
    doc.add_paragraph('Racial Diversity: {{ diverse_percentage }}%')

def create_board_analysis_template(doc):
    title = doc.add_heading('Board Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Client: {{ client_name }}')
    doc.add_paragraph('Board: {{ board_name }}')
    doc.add_paragraph('Ticker: {{ company_ticker }}')
    doc.add_paragraph('Date: {{ report_date }}')
    doc.add_paragraph()
    
    doc.add_heading('Board Composition', 1)
    doc.add_paragraph('Total Members: {{ total_members }}')
    doc.add_paragraph()
    
    doc.add_heading('Diversity Metrics', 1)
    doc.add_paragraph('Female: {{ female_count }} ({{ female_percentage }}%)')
    doc.add_paragraph('Diverse: {{ diverse_count }} ({{ diverse_percentage }}%)')
    doc.add_paragraph('S&P 500: {{ sp500_female }}% female, {{ sp500_diverse }}% diverse')

def create_simple_template(doc, doc_type):
    title = doc.add_heading(f'{doc_type.replace("_", " ").title()} Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Client: {{ client_name }}')
    doc.add_paragraph('Date: {{ report_date }}')
    doc.add_paragraph()
    doc.add_paragraph(f'Report content for {doc_type}')
